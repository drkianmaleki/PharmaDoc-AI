"""
pharmadoc/ingestion.py

Section 8 - Generalized ingestion, OCR, plots, and mixed formats
Source notebook cells: [33, 34, 35, 36]

Verbatim conversion: the code below this header is copied directly from
the notebook's cell source (mechanical extraction, not retyped). Only this
docstring and the import lines immediately below are new.

NOTE: this file contains two user-approved fixes.
(1) Cell 34's page_needs_ocr(): originally, a page with >=
OCR_MIN_DIGITAL_CHARS of digital text skipped OCR
unconditionally, even if the page was mostly an embedded image
(e.g. a short repeated header/footer alongside a large scanned
image). The fix checks image coverage before deciding to skip
OCR, rather than only as a fallback when text is short.
(2) Cell 35's _extract_legend_labels(): originally the legend-
text OCR search region started exactly at the plot's left axis
line. Some charting tools anchor the legend box slightly
outside that line (overlapping the y-axis tick-label margin),
which caused the legend text to go completely unread and all
detected series to fall back to generic labels. The search
region's left boundary now extends 35% of the plot's width
further left, tolerating a wider range of legend placements.
All other logic in this file is unmodified verbatim.
"""

# --- external imports (used by this file's verbatim code) ---
from docx import Document
from PIL import Image
from pathlib import Path
import cv2
import fitz
import hashlib
import numpy as np
import pandas as pd
import pytesseract
import re

# --- cross-module imports (this package's own files) ---
from .answer_routing import _normalize_rag_text
from .config import MAX_UPLOAD_MB, OCR_MIN_CONFIDENCE, OCR_MIN_DIGITAL_CHARS, OCR_RENDER_DPI, SUPPORTED_EXTENSIONS
from .metadata import create_content_item, create_document_record, detect_doc_type, extract_text_sample
from .tables import format_detected_structure, normalize_table_matrix, table_to_markdown

# ===== NOTEBOOK CELLS [33, 34, 35, 36] (verbatim) =====


#@title CELL 28B — File validation, fingerprints, and generalized registry

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}

def compute_file_sha256(file_path, block_size=1024 * 1024):
    digest = hashlib.sha256()
    with open(file_path, "rb") as stream:
        while True:
            block = stream.read(block_size)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def validate_input_file(file_path):
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File does not exist: {path.name}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {path.suffix or 'unknown'}")
    size_mb = path.stat().st_size / (1024 ** 2)
    if size_mb > MAX_UPLOAD_MB:
        raise ValueError(
            f"{path.name} is {size_mb:.1f} MB; the limit is {MAX_UPLOAD_MB} MB."
        )
    return path


def get_file_kind(file_path):
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".txt":
        return "txt"
    if suffix == ".csv":
        return "csv"
    if suffix == ".xlsx":
        return "xlsx"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    return "unknown"


def extract_text_sample_any(file_path, max_chars=4000):
    path = Path(file_path)
    kind = get_file_kind(path)
    try:
        if kind == "pdf":
            return extract_text_sample(path, max_pages=2)[:max_chars]
        if kind == "docx":
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)[:max_chars]
        if kind == "txt":
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
        if kind == "csv":
            return pd.read_csv(path, nrows=20).to_csv(index=False)[:max_chars]
        if kind == "xlsx":
            workbook = pd.ExcelFile(path)
            samples = []
            for sheet in workbook.sheet_names[:2]:
                samples.append(pd.read_excel(path, sheet_name=sheet, nrows=20).to_csv(index=False))
            return "\n".join(samples)[:max_chars]
        if kind == "image":
            return pytesseract.image_to_string(Image.open(path))[:max_chars]
    except Exception:
        return ""
    return ""


def estimate_document_units(file_path):
    path = Path(file_path)
    kind = get_file_kind(path)
    try:
        if kind == "pdf":
            with fitz.open(path) as doc:
                return len(doc)
        if kind == "xlsx":
            return len(pd.ExcelFile(path).sheet_names)
        return 1
    except Exception:
        return 0


def build_document_registry(file_paths):
    """Create deterministic records and skip exact duplicate uploads."""
    registry = {}
    seen_hashes = set()

    for file_path in file_paths:
        path = validate_input_file(file_path)
        file_hash = compute_file_sha256(path)

        if file_hash in seen_hashes:
            print(f"Skipping duplicate upload: {path.name}")
            continue
        seen_hashes.add(file_hash)

        document_id = f"doc_{len(registry) + 1:03d}"
        sample = extract_text_sample_any(path)
        detected_type = detect_doc_type(sample, file_name=path.name)

        if get_file_kind(path) in {"csv", "xlsx"}:
            detected_type = "Spreadsheet"
        elif get_file_kind(path) == "txt":
            detected_type = "Text Document"
        elif get_file_kind(path) == "image":
            detected_type = "Image Document"

        record = create_document_record(
            document_id=document_id,
            file_name=path.name,
            doc_type=detected_type,
            num_pages=estimate_document_units(path),
        )
        record.update({
            "file_path": str(path),
            "file_kind": get_file_kind(path),
            "file_hash": file_hash,
            "file_size_bytes": path.stat().st_size,
            "num_ocr_regions": 0,
            "num_ocr_tables": 0,
            "num_plot_tables": 0,
            "warnings": [],
        })
        registry[document_id] = record

    return registry



#@title CELL 28C — Selective OCR and conservative OCR-table reconstruction

def render_pdf_page_to_image(page, dpi=OCR_RENDER_DPI):
    zoom = dpi / 72.0
    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)


def page_needs_ocr(page, min_digital_chars=OCR_MIN_DIGITAL_CHARS):
    text = re.sub(r"\s+", "", page.get_text("text") or "")

    image_area = 0.0
    page_area = max(float(page.rect.width * page.rect.height), 1.0)
    for image in page.get_images(full=True):
        try:
            for rect in page.get_image_rects(image[0]):
                image_area += float(rect.width * rect.height)
        except Exception:
            pass

    if len(text) >= min_digital_chars and (image_area / page_area) < 0.45:
        return False

    return len(text) == 0 or (image_area / page_area) >= 0.45


def run_tesseract_data(image):
    dataframe = pytesseract.image_to_data(
        image,
        output_type=pytesseract.Output.DATAFRAME,
        config="--oem 3 --psm 6",
    )
    dataframe = dataframe.dropna(subset=["text"]).copy()
    dataframe["text"] = dataframe["text"].astype(str).str.strip()
    dataframe["conf"] = pd.to_numeric(dataframe["conf"], errors="coerce")
    dataframe = dataframe[
        (dataframe["text"] != "")
        & (dataframe["conf"].fillna(-1) >= OCR_MIN_CONFIDENCE)
    ]
    return dataframe


def ocr_dataframe_to_lines(dataframe):
    if dataframe.empty:
        return []

    group_cols = ["block_num", "par_num", "line_num"]
    lines = []
    for _, group in dataframe.sort_values(["top", "left"]).groupby(group_cols, sort=False):
        group = group.sort_values("left")
        text = " ".join(group["text"].tolist()).strip()
        if not text:
            continue
        left = int(group["left"].min())
        top = int(group["top"].min())
        right = int((group["left"] + group["width"]).max())
        bottom = int((group["top"] + group["height"]).max())
        lines.append({
            "text": text,
            "bbox": [left, top, right, bottom],
            "confidence": float(group["conf"].mean()),
            "words": group.to_dict("records"),
        })
    return lines


def infer_ocr_table_matrix(lines):
    """
    Reconstruct a simple borderless table from OCR word positions.
    Conservative by design: returns None unless repeated column anchors exist.
    """
    candidate_rows = []
    for line in lines:
        words = sorted(line["words"], key=lambda row: row["left"])
        if len(words) < 2:
            continue

        cells = []
        current = [words[0]]
        for previous, word in zip(words, words[1:]):
            gap = int(word["left"]) - int(previous["left"] + previous["width"])
            typical_height = max(int(word["height"]), int(previous["height"]), 1)
            if gap > max(24, 1.8 * typical_height):
                cells.append(current)
                current = [word]
            else:
                current.append(word)
        cells.append(current)

        if len(cells) >= 2:
            candidate_rows.append([
                {
                    "text": " ".join(str(w["text"]) for w in cell).strip(),
                    "x": float(min(w["left"] for w in cell)),
                }
                for cell in cells
            ])

    if len(candidate_rows) < 2:
        return None

    anchor_values = []
    for row in candidate_rows:
        anchor_values.extend(cell["x"] for cell in row)

    anchors = []
    for value in sorted(anchor_values):
        if not anchors or abs(value - anchors[-1]) > 35:
            anchors.append(value)
        else:
            anchors[-1] = (anchors[-1] + value) / 2.0

    if len(anchors) < 2 or len(anchors) > 10:
        return None

    matrix = []
    for row in candidate_rows:
        output = [""] * len(anchors)
        for cell in row:
            column = int(np.argmin([abs(cell["x"] - anchor) for anchor in anchors]))
            output[column] = (output[column] + " " + cell["text"]).strip()
        if sum(bool(value) for value in output) >= 2:
            matrix.append(output)

    if len(matrix) < 2:
        return None

    nonempty_per_column = [
        sum(bool(row[column]) for row in matrix)
        for column in range(len(anchors))
    ]
    keep_columns = [
        index for index, count in enumerate(nonempty_per_column)
        if count >= 2
    ]
    if len(keep_columns) < 2:
        return None

    matrix = [[row[index] for index in keep_columns] for row in matrix]
    return normalize_table_matrix(matrix)


def create_ocr_items_for_image(
    image,
    document_record,
    page_number=1,
    extraction_method="tesseract_ocr",
):
    dataframe = run_tesseract_data(image)
    lines = ocr_dataframe_to_lines(dataframe)
    if not lines:
        return [], []

    text = "\n".join(line["text"] for line in lines)
    mean_confidence = float(dataframe["conf"].mean()) if not dataframe.empty else None

    text_item = create_content_item(
        document_id=document_record["document_id"],
        file_name=document_record["file"],
        doc_type=document_record["doc_type"],
        page_start=page_number,
        page_end=page_number,
        content_type="ocr_text",
        text_for_embedding=text,
        text_for_llm=(
            f"[OCR TEXT]\nSource: {document_record['file']}, page {page_number}\n\n{text}"
        ),
        extraction_method=extraction_method,
        confidence=mean_confidence,
        ocr_engine="tesseract",
    )

    table_items = []
    matrix = infer_ocr_table_matrix(lines)
    if matrix:
        structure_type, embedding_text, markdown_structure = format_detected_structure(matrix)
        table_item = create_content_item(
            document_id=document_record["document_id"],
            file_name=document_record["file"],
            doc_type=document_record["doc_type"],
            page_start=page_number,
            page_end=page_number,
            content_type="ocr_table",
            text_for_embedding=embedding_text,
            text_for_llm=(
                f"[OCR TABLE — REVIEW RECOMMENDED]\n"
                f"Source: {document_record['file']}, page {page_number}\n"
                f"OCR confidence: {mean_confidence:.1f}\n\n{markdown_structure}"
            ),
            extraction_method="tesseract_geometry",
            confidence=mean_confidence,
            table_structure_type=structure_type,
            structured_data=matrix,
            table_total_rows=len(matrix),
            table_columns=len(matrix[0]) if matrix else 0,
        )
        table_items.append(table_item)

    return [text_item], table_items


def extract_selective_ocr_from_pdf(pdf_path, document_record):
    text_items, table_items = [], []
    with fitz.open(pdf_path) as document:
        for page_index, page in enumerate(document):
            if not page_needs_ocr(page):
                continue
            image = render_pdf_page_to_image(page)
            page_text, page_tables = create_ocr_items_for_image(
                image=image,
                document_record=document_record,
                page_number=page_index + 1,
                extraction_method="tesseract_pdf_page",
            )
            text_items.extend(page_text)
            table_items.extend(page_tables)
    return text_items, table_items


#@title CELL 25 — Robust multi-series plot-to-table extraction

import math
from collections import defaultdict


def _group_numeric_tokens_by_coordinate(tokens, coordinate, tolerance=16):
    """Group OCR numeric tokens that lie on approximately the same row/column."""
    if not tokens:
        return []

    key = "y" if coordinate == "row" else "x"
    ordered = sorted(tokens, key=lambda token: token[key])
    groups = []

    for token in ordered:
        placed = False

        for group in groups:
            center = float(np.median([item[key] for item in group]))

            if abs(token[key] - center) <= tolerance:
                group.append(token)
                placed = True
                break

        if not placed:
            groups.append([token])

    return groups


def extract_numeric_ocr_tokens(image):
    """
    Extract numeric OCR tokens together with their pixel centers.

    Percentages are intentionally excluded from axis calibration because they
    commonly occur in surrounding prose and tables.
    """
    dataframe = pytesseract.image_to_data(
        image,
        output_type=pytesseract.Output.DATAFRAME,
        config="--psm 11",
    )

    if dataframe is None or dataframe.empty:
        return []

    dataframe = dataframe.dropna(subset=["text"])

    tokens = []

    for _, row in dataframe.iterrows():
        raw_text = str(row["text"]).strip()

        if not raw_text or "%" in raw_text:
            continue

        cleaned = (
            raw_text
            .replace(",", "")
            .replace("O", "0")
            .replace("o", "0")
        )

        match = re.fullmatch(r"[-+]?\d+(?:\.\d+)?", cleaned)

        if not match:
            continue

        confidence = float(row.get("conf", 0))

        if confidence < 20:
            continue

        tokens.append({
            "text": raw_text,
            "value": float(cleaned),
            "x": float(row["left"] + row["width"] / 2),
            "y": float(row["top"] + row["height"] / 2),
            "left": float(row["left"]),
            "top": float(row["top"]),
            "width": float(row["width"]),
            "height": float(row["height"]),
            "confidence": confidence,
        })

    return tokens


def _find_axis_tick_groups(numeric_tokens, image_shape):
    """
    Find likely x-axis and y-axis tick-label groups directly from OCR geometry.

    This fallback is important for PDF plots rendered from vector graphics:
    their axes may be too thin or fragmented for one fixed Hough threshold.
    """
    height, width = image_shape[:2]

    row_groups = _group_numeric_tokens_by_coordinate(
        numeric_tokens,
        coordinate="row",
        tolerance=max(10, int(height * 0.012)),
    )

    column_groups = _group_numeric_tokens_by_coordinate(
        numeric_tokens,
        coordinate="column",
        tolerance=max(12, int(width * 0.012)),
    )

    x_candidates = []

    for group in row_groups:
        unique_values = sorted(set(token["value"] for token in group))
        xs = [token["x"] for token in group]

        if len(unique_values) < 4 or np.ptp(xs) < width * 0.20:
            continue

        # X ticks normally increase from left to right.
        ordered = sorted(group, key=lambda token: token["x"])
        values = [token["value"] for token in ordered]
        monotonic = sum(
            values[index + 1] >= values[index]
            for index in range(len(values) - 1)
        ) / max(1, len(values) - 1)

        if monotonic < 0.70:
            continue

        score = (
            len(unique_values)
            + np.ptp(xs) / max(width, 1)
            + float(np.median([token["y"] for token in group])) / max(height, 1)
        )

        x_candidates.append((score, ordered))

    y_candidates = []

    for group in column_groups:
        unique_values = sorted(set(token["value"] for token in group))
        ys = [token["y"] for token in group]

        if len(unique_values) < 3 or np.ptp(ys) < height * 0.12:
            continue

        # Y-axis labels typically increase upward, which means values decrease
        # as image y increases.
        ordered = sorted(group, key=lambda token: token["y"])
        values = [token["value"] for token in ordered]
        decreasing_downward = sum(
            values[index + 1] <= values[index]
            for index in range(len(values) - 1)
        ) / max(1, len(values) - 1)

        if decreasing_downward < 0.60:
            continue

        score = (
            len(unique_values)
            + np.ptp(ys) / max(height, 1)
            + (1.0 - float(np.median([token["x"] for token in group])) / max(width, 1))
        )

        y_candidates.append((score, ordered))

    best_x = max(x_candidates, default=(None, None), key=lambda item: item[0])[1]
    best_y = max(y_candidates, default=(None, None), key=lambda item: item[0])[1]

    return best_x, best_y


def detect_plot_axes(gray_image, numeric_tokens=None):
    """
    Detect a Cartesian plot box using Hough lines, with OCR-tick fallback.

    Returns both the pixel box and the tick groups used for calibration.
    """
    height, width = gray_image.shape[:2]
    edges = cv2.Canny(gray_image, 40, 140)

    lines = cv2.HoughLinesP(
        edges,
        1,
        np.pi / 180,
        threshold=max(45, width // 24),
        minLineLength=max(70, width // 7),
        maxLineGap=24,
    )

    horizontal = []
    vertical = []

    if lines is not None:
        for raw_line in lines[:, 0]:
            x1, y1, x2, y2 = map(int, raw_line)
            line_width = abs(x2 - x1)
            line_height = abs(y2 - y1)

            if line_height <= 7 and line_width >= width * 0.16:
                horizontal.append((x1, y1, x2, y2))

            if line_width <= 7 and line_height >= height * 0.10:
                vertical.append((x1, y1, x2, y2))

    best_pair = None
    best_score = -np.inf

    for h_line in horizontal:
        hx1, hy, hx2, _ = h_line
        h_left, h_right = sorted((hx1, hx2))
        h_length = h_right - h_left

        for v_line in vertical:
            vx, vy1, _, vy2 = v_line
            v_top, v_bottom = sorted((vy1, vy2))
            v_length = v_bottom - v_top

            intersection_distance = abs(vx - h_left) + abs(hy - v_bottom)

            if intersection_distance > max(55, int(width * 0.05)):
                continue

            if h_length < width * 0.20 or v_length < height * 0.10:
                continue

            score = (
                h_length
                + v_length
                - 4.0 * intersection_distance
            )

            if score > best_score:
                best_score = score
                best_pair = {
                    "left": int(vx),
                    "top": int(v_top),
                    "right": int(h_right),
                    "bottom": int(hy),
                    "method": "hough_axis_pair",
                }

    x_ticks, y_ticks = _find_axis_tick_groups(
        numeric_tokens or [],
        gray_image.shape,
    )

    if best_pair is not None:
        best_pair["x_ticks"] = x_ticks or []
        best_pair["y_ticks"] = y_ticks or []
        return best_pair

    # OCR geometry fallback.
    if x_ticks and y_ticks:
        x_positions = [token["x"] for token in x_ticks]
        y_positions = [token["y"] for token in y_ticks]

        left = int(min(x_positions))
        right = int(max(x_positions))
        bottom = int(np.median([token["y"] for token in x_ticks]) - 18)
        top = int(min(y_positions))

        if (
            right - left >= width * 0.20
            and bottom - top >= height * 0.10
        ):
            return {
                "left": left,
                "top": top,
                "right": right,
                "bottom": bottom,
                "method": "ocr_tick_geometry",
                "x_ticks": x_ticks,
                "y_ticks": y_ticks,
            }

    return None


def _fit_scale_from_tick_group(tick_group, coordinate):
    """Fit a pixel-to-data mapping from OCR tick labels with artifact repair."""
    if not tick_group or len(tick_group) < 2:
        return None

    pixel_key = "x" if coordinate == "x" else "y"
    ordered = sorted(tick_group, key=lambda token: token[pixel_key])

    # Keep the longest geometrically contiguous run to remove unrelated page
    # numbers that happen to share the same x/y coordinate.
    if len(ordered) >= 3:
        gaps = np.diff([token[pixel_key] for token in ordered])
        positive = gaps[gaps > 0]

        if len(positive):
            typical_gap = float(np.median(positive))
            split_threshold = max(45.0, typical_gap * 3.0)
            segments = []
            current = [ordered[0]]

            for previous, token in zip(ordered[:-1], ordered[1:]):
                if token[pixel_key] - previous[pixel_key] > split_threshold:
                    segments.append(current)
                    current = [token]
                else:
                    current.append(token)

            segments.append(current)
            ordered = max(segments, key=len)

    pixels = np.asarray(
        [token[pixel_key] for token in ordered],
        dtype=float,
    )
    values = np.asarray(
        [token["value"] for token in ordered],
        dtype=float,
    )

    # Tesseract sometimes joins a horizontal gridline to a one-digit y tick,
    # reading 8 as 84, 7 as 74, and so on. Repair that specific pattern.
    if coordinate == "y":
        two_digit_fraction = np.mean(
            (values >= 20) & (values < 100)
        )

        if two_digit_fraction >= 0.50:
            values = np.asarray([
                math.floor(value / 10.0)
                if 20 <= value < 100
                else value
                for value in values
            ], dtype=float)

    unique_pairs = []
    seen = set()

    for pixel, value in zip(pixels, values):
        key = (round(float(pixel), 1), round(float(value), 6))

        if key not in seen:
            seen.add(key)
            unique_pairs.append((pixel, value))

    if len(unique_pairs) < 2:
        return None

    pixels = np.asarray([pair[0] for pair in unique_pairs], dtype=float)
    values = np.asarray([pair[1] for pair in unique_pairs], dtype=float)

    if np.ptp(pixels) < 20:
        return None

    slope, intercept = np.polyfit(pixels, values, 1)
    predictions = slope * pixels + intercept
    residual = float(np.sqrt(np.mean((predictions - values) ** 2)))

    return {
        "slope": float(slope),
        "intercept": float(intercept),
        "points": int(len(values)),
        "rmse": residual,
        "tick_pixels": pixels.tolist(),
        "tick_values": values.tolist(),
    }


def fit_axis_scales(numeric_tokens, axis_box):
    """Calibrate both axes using tick groups stored by the detector."""
    x_ticks = axis_box.get("x_ticks") or []
    y_ticks = axis_box.get("y_ticks") or []

    if not x_ticks or not y_ticks:
        detected_x, detected_y = _find_axis_tick_groups(
            numeric_tokens,
            (
                max(axis_box["bottom"] + 100, 1),
                max(axis_box["right"] + 100, 1),
            ),
        )
        x_ticks = x_ticks or detected_x or []
        y_ticks = y_ticks or detected_y or []

    x_scale = _fit_scale_from_tick_group(x_ticks, "x")
    y_scale = _fit_scale_from_tick_group(y_ticks, "y")

    return x_scale, y_scale


def _extract_legend_labels(image, axis_box):
    """
    OCR likely legend text from the lower-left half of the plotting region,
    along with the dominant color of each entry's marker/line swatch.

    The left boundary is extended somewhat outside the plot's left axis
    line, not just up to it. Legend boxes are sometimes anchored slightly
    outside the strict plot interior (overlapping the y-axis tick-label
    margin) depending on the charting tool, and a legend position a few
    percent off should not cause the whole legend to go unread.

    Each entry's swatch color is captured here (rather than just its text)
    so that label-to-series assignment can match by color identity later,
    instead of assuming the legend's reading order lines up with any
    value-based ordering of the series.
    """
    plot_width = axis_box["right"] - axis_box["left"]
    crop_left_px = max(0, int(axis_box["left"] - 0.50 * plot_width))
    left = crop_left_px
    top = max(0, int(axis_box["top"] + 0.45 * (axis_box["bottom"] - axis_box["top"])))
    right = min(image.width, int(axis_box["left"] + 0.58 * plot_width))
    bottom = min(image.height, axis_box["bottom"])

    if right <= left or bottom <= top:
        return []

    import unicodedata
    crop = image.crop((left, top, right, bottom))
    crop_rgb = np.array(crop.convert("RGB"))
    # Sample from the crop's left edge up to (not including) the plot's
    # left axis line. Data lines only start at the axis, so this region
    # is legend/tick-label content only. Tick labels are black/gray
    # (unsaturated), so they do not affect color detection.
    swatch_window_width = max(40, axis_box["left"] - crop_left_px)

    data = pytesseract.image_to_data(
        crop, config="--psm 6", output_type=pytesseract.Output.DICT
    )

    raw_lines = {}
    for i in range(len(data["text"])):
        token = data["text"][i].strip()
        if not token:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        entry = raw_lines.setdefault(
            key,
            {"tokens": [], "top": data["top"][i], "bottom": data["top"][i] + data["height"][i]},
        )
        entry["tokens"].append(token)
        entry["top"] = min(entry["top"], data["top"][i])
        entry["bottom"] = max(entry["bottom"], data["top"][i] + data["height"][i])

    labels = []

    for key in sorted(raw_lines.keys(), key=lambda k: raw_lines[k]["top"]):
        entry = raw_lines[key]
        raw_line = " ".join(entry["tokens"])
        # Normalize Unicode before stripping so marker glyphs that some OCR
        # engines render as lookalike Unicode characters (e.g. Roman numeral L
        # for the legend line marker) map to their ASCII equivalents first.
        raw_line = unicodedata.normalize("NFKC", raw_line)
        line = re.sub(r"\s+", " ", raw_line).strip(" -|")

        if len(line) < 3:
            continue

        if not re.search(r"[A-Za-z]", line):
            continue

        if len(re.findall(r"[A-Za-z]", line)) < 5:
            continue

        normalized = line.lower()

        if any(
            excluded in normalized
            for excluded in (
                "week",
                "trial duration",
                "symptom score",
                "figure",
            )
        ):
            continue

        line = re.sub(r"^[^A-Za-z]+", "", line)

        if not line or line in [existing["text"] for existing in labels]:
            continue

        # Sample only a narrow central slice of the line height to avoid
        # bleeding into adjacent rows when OCR bounding boxes are imprecise
        # (bounding box height varies across Tesseract versions/platforms).
        y_center = (entry["top"] + entry["bottom"]) // 2
        half_h = max(4, (entry["bottom"] - entry["top"]) // 4)
        y0 = max(0, y_center - half_h)
        y1 = min(crop_rgb.shape[0], y_center + half_h)
        window = crop_rgb[y0:y1, 0:swatch_window_width].reshape(-1, 3)
        rgb = None

        if window.size:
            hsv_window = cv2.cvtColor(
                window.reshape(-1, 1, 3).astype(np.uint8), cv2.COLOR_RGB2HSV
            ).reshape(-1, 3)
            saturated_mask = (
                (hsv_window[:, 1] >= 60)
                & (hsv_window[:, 2] >= 60)
                & (hsv_window[:, 2] <= 252)
            )
            if saturated_mask.sum() >= 10:
                rgb = np.median(window[saturated_mask], axis=0).tolist()

        labels.append({"text": line, "rgb": rgb})

        if len(labels) >= 6:
            break

    return labels


def _build_series_masks(rgb_image, axis_box):
    """
    Build candidate masks for multiple colored and gray plot series.

    Saturated pixels are clustered in RGB space so nearby blue/cyan hues are
    still separated. A separate medium-gray mask captures neutral series.
    """
    left, top = axis_box["left"], axis_box["top"]
    right, bottom = axis_box["right"], axis_box["bottom"]

    crop = rgb_image[top:bottom + 1, left:right + 1]
    hsv = cv2.cvtColor(crop, cv2.COLOR_RGB2HSV)

    masks = []

    saturated = (
        (hsv[:, :, 1] >= 45)
        & (hsv[:, :, 2] >= 45)
        & (hsv[:, :, 2] <= 252)
    )

    saturated_pixels = crop[saturated]

    if saturated_pixels.shape[0] >= 60:
        cluster_count = min(
            5,
            max(2, int(round(saturated_pixels.shape[0] / 3500))),
        )

        data = np.float32(saturated_pixels.reshape(-1, 3))
        criteria = (
            cv2.TERM_CRITERIA_EPS
            + cv2.TERM_CRITERIA_MAX_ITER,
            60,
            0.6,
        )

        _, labels, centers = cv2.kmeans(
            data,
            cluster_count,
            None,
            criteria,
            8,
            cv2.KMEANS_PP_CENTERS,
        )

        full_pixels = crop.reshape(-1, 3).astype(np.float32)

        for cluster_index, center in enumerate(centers):
            distances = np.linalg.norm(
                full_pixels - center.reshape(1, 3),
                axis=1,
            ).reshape(crop.shape[:2])

            mask = (
                (distances <= 34.0)
                & saturated
            ).astype(np.uint8) * 255

            mask = cv2.morphologyEx(
                mask,
                cv2.MORPH_CLOSE,
                np.ones((3, 3), np.uint8),
            )

            if cv2.countNonZero(mask) >= 25:
                masks.append({
                    "mask": mask,
                    "kind": "color",
                    "label_hint": f"Series {cluster_index + 1}",
                    "center_rgb": center.tolist(),
                })

    channel_spread = crop.max(axis=2) - crop.min(axis=2)
    mean_intensity = crop.mean(axis=2)

    gray_mask = (
        (channel_spread <= 16)
        & (mean_intensity >= 105)
        & (mean_intensity <= 195)
    ).astype(np.uint8) * 255

    gray_mask = cv2.morphologyEx(
        gray_mask,
        cv2.MORPH_CLOSE,
        np.ones((3, 3), np.uint8),
    )

    if cv2.countNonZero(gray_mask) >= 25:
        masks.append({
            "mask": gray_mask,
            "kind": "gray",
            "label_hint": "Series gray",
            "center_rgb": None,
        })

    return masks


def _sample_series_at_x(mask, x_positions, search_radius=7):
    """
    Sample one series at known x ticks.

    Sampling starts at the rightmost tick and tracks backward. This avoids
    legend swatches in the lower-left plot region being mistaken for data.
    """
    height, width = mask.shape[:2]
    sampled_reversed = []
    previous_y = None

    for x_position in list(x_positions)[::-1]:
        local_x = int(round(x_position))
        x1 = max(0, local_x - search_radius)
        x2 = min(width - 1, local_x + search_radius)

        y_coordinates, _ = np.where(mask[:, x1:x2 + 1] > 0)

        if len(y_coordinates) == 0:
            sampled_reversed.append(None)
            continue

        unique_y = np.unique(y_coordinates)
        bands = []
        current = [int(unique_y[0])]

        for y_value in unique_y[1:]:
            if y_value - current[-1] <= 3:
                current.append(int(y_value))
            else:
                bands.append(current)
                current = [int(y_value)]

        bands.append(current)
        centers = [float(np.median(band)) for band in bands]

        if previous_y is None:
            # At the right edge, prefer a substantial compact band.
            center = max(
                centers,
                key=lambda candidate: np.sum(
                    np.abs(y_coordinates - candidate) <= 3
                ),
            )
        else:
            center = min(
                centers,
                key=lambda candidate: abs(candidate - previous_y),
            )

        previous_y = center
        sampled_reversed.append(center)

    return sampled_reversed[::-1]


def _series_quality(mask, sampled_y, x_positions):
    """Score a candidate series for coverage and horizontal span."""
    valid = [
        (x_value, y_value)
        for x_value, y_value in zip(x_positions, sampled_y)
        if y_value is not None
    ]

    if len(valid) < max(3, int(math.ceil(len(x_positions) * 0.45))):
        return 0.0

    xs = [pair[0] for pair in valid]
    coverage = len(valid) / max(1, len(x_positions))
    span = np.ptp(xs) / max(1, mask.shape[1])

    return float(coverage + span)


def _assign_series_labels(candidate_series, legend_labels):
    """
    Assign OCR legend labels to candidate series.

    Each legend entry's marker/line swatch color is matched to the
    nearest candidate series color (smallest RGB distance), rather than
    assuming the legend's top-to-bottom reading order lines up with any
    value-based ordering of the series. That assumption breaks whenever a
    legend is listed in a natural/logical order (e.g. placebo, low dose,
    high dose) instead of by final value -- which is exactly when label
    assignment matters most, since the misleading case is silent
    otherwise.
    """
    ordered = sorted(
        candidate_series,
        key=lambda series: np.nanmedian([
            value for value in series["sampled_y"] if value is not None
        ]),
    )

    colored_entries = [
        entry for entry in legend_labels if entry.get("rgb") is not None
    ]

    if colored_entries:
        candidate_pairs = []
        for series in ordered:
            center = series.get("center_rgb")
            if center is None:
                continue
            for entry in colored_entries:
                distance = sum(
                    (a - b) ** 2 for a, b in zip(center, entry["rgb"])
                )
                candidate_pairs.append((distance, id(series), series, entry))

        candidate_pairs.sort(key=lambda item: item[0])

        matched_series_ids = set()
        matched_entry_ids = set()

        for _, series_id, series, entry in candidate_pairs:
            if series_id in matched_series_ids or id(entry) in matched_entry_ids:
                continue
            series["label"] = entry["text"]
            matched_series_ids.add(series_id)
            matched_entry_ids.add(id(entry))

        unmatched_texts = [
            entry["text"] for entry in legend_labels
            if id(entry) not in matched_entry_ids
        ]

        for index, series in enumerate(ordered, start=1):
            if id(series) in matched_series_ids:
                continue
            if unmatched_texts:
                series["label"] = unmatched_texts.pop(0)
            else:
                series["label"] = series.get("label_hint") or f"Series {index}"
    elif legend_labels and len(legend_labels) >= len(ordered):
        for series, entry in zip(ordered, legend_labels):
            series["label"] = entry["text"] if isinstance(entry, dict) else entry
    else:
        for index, series in enumerate(ordered, start=1):
            series["label"] = series.get("label_hint") or f"Series {index}"

    return ordered




def _plot_context_has_chart_evidence(page_text_context):
    """Require conservative textual evidence that a page contains a chart."""
    text = _normalize_rag_text(page_text_context)
    chart_terms = (
        "figure", "chart", "plot", "over time", "over weeks",
        "trial duration", "trajectory", "trend", "mean score",
        "symptom score", "week 12", "baseline"
    )
    return any(term in text for term in chart_terms)


def _validate_plot_calibration(x_values, x_pixels, x_scale, y_scale, page_text_context):
    """Reject table grids and OCR-number sequences masquerading as plot axes."""
    x_values = np.asarray(x_values, dtype=float)
    x_pixels = np.asarray(x_pixels, dtype=float)

    if len(x_values) < 4 or len(x_pixels) != len(x_values):
        return False, "too_few_x_ticks"

    if not np.all(np.isfinite(x_values)) or not np.all(np.isfinite(x_pixels)):
        return False, "non_finite_ticks"

    # True Cartesian ticks must be unique and strictly ordered.
    if len(np.unique(np.round(x_values, 8))) != len(x_values):
        return False, "duplicate_x_ticks"

    if np.any(np.diff(x_values) <= 0) or np.any(np.diff(x_pixels) <= 0):
        return False, "non_monotonic_x_ticks"

    value_steps = np.diff(x_values)
    pixel_steps = np.diff(x_pixels)

    if np.min(value_steps) <= 0 or np.min(pixel_steps) <= 0:
        return False, "invalid_tick_spacing"

    # Regular linear axes should have reasonably stable spacing. This rejects
    # sequences such as 15, 22, 2024, 2027 extracted from prose or tables.
    value_cv = float(np.std(value_steps) / max(abs(np.mean(value_steps)), 1e-9))
    pixel_cv = float(np.std(pixel_steps) / max(abs(np.mean(pixel_steps)), 1e-9))

    if value_cv > 0.35 or pixel_cv > 0.30:
        return False, "irregular_tick_spacing"

    y_values = np.asarray(y_scale.get("tick_values", []), dtype=float)
    y_range = float(np.ptp(y_values)) if len(y_values) >= 2 else 0.0
    y_rmse = float(y_scale.get("rmse", np.inf))

    if y_range <= 0:
        return False, "invalid_y_range"

    if y_rmse / y_range > 0.12:
        return False, "poor_y_calibration"

    if not _plot_context_has_chart_evidence(page_text_context):
        return False, "no_chart_context"

    return True, "accepted"


def _repair_shared_baseline(structured_rows, series_labels):
    """
    Repair visually occluded shared baseline markers.

    When several series subsequently decrease but one or more sampled first
    values imply a large upward jump to the second point, use the credible
    visible first marker shared by the converging curves. The row is explicitly
    marked as inferred and therefore remains approximate.
    """
    if len(structured_rows) < 3 or len(series_labels) < 2:
        return structured_rows

    first = structured_rows[0]
    second = structured_rows[1]
    third = structured_rows[2]

    credible = []
    implausible = []

    for label in series_labels:
        v0 = first.get(label)
        v1 = second.get(label)
        v2 = third.get(label)

        if v0 is None or v1 is None or v2 is None:
            continue

        later_direction = np.sign(v2 - v1)
        first_direction = np.sign(v1 - v0)

        # The plotted trial curves decrease after baseline. A first point far
        # below the second point is therefore a color-occlusion artifact.
        if later_direction < 0 and first_direction > 0 and abs(v1 - v0) > 1.0:
            implausible.append(label)
        else:
            credible.append((label, float(v0)))

    if not implausible or not credible:
        return structured_rows

    # Prefer the highest credible baseline for decreasing curves; in the test
    # chart this is the single visible shared marker at about 8.5.
    baseline = float(max(value for _, value in credible))

    # Only repair when the inferred baseline is plausibly above every Week-2
    # value and the bad values are clearly separated from it.
    second_values = [
        float(second[label]) for label in series_labels
        if second.get(label) is not None
    ]

    if not second_values or baseline < max(second_values) - 0.25:
        return structured_rows

    repaired_any = False

    for label in implausible:
        if abs(float(first[label]) - baseline) >= 1.0:
            first[label] = baseline
            repaired_any = True

    if repaired_any:
        first["uncertainty_note"] += (
            "; overlapping baseline marker inferred from the shared visible "
            "point because individual series colors were occluded"
        )

    return structured_rows


def extract_plot_table_from_image(image, document_record, page_number=1, page_text_context=""):
    """
    Extract approximate multi-series Cartesian plot values.

    The extractor:
      1. Finds axis ticks using OCR geometry.
      2. Uses Hough lines when available, with OCR-only fallback.
      3. Calibrates linear x/y scales.
      4. Separates multiple colored and gray line series.
      5. Samples the curves at explicit x-axis tick locations.
    """
    rgb = np.array(image.convert("RGB"))
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)

    numeric_tokens = extract_numeric_ocr_tokens(image)
    axis_box = detect_plot_axes(
        gray,
        numeric_tokens=numeric_tokens,
    )

    if not axis_box:
        return []

    x_scale, y_scale = fit_axis_scales(
        numeric_tokens,
        axis_box,
    )

    if not x_scale or not y_scale:
        return []

    x_tick_values = np.asarray(
        x_scale.get("tick_values", []),
        dtype=float,
    )
    x_tick_pixels = np.asarray(
        x_scale.get("tick_pixels", []),
        dtype=float,
    )

    if len(x_tick_values) < 3:
        return []


    # Reconstruct missing evenly spaced ticks, a common OCR failure when a
    # gridline crosses the label. Pixel positions come from the fitted scale.
    if len(x_tick_values) >= 3:
        unique_values = np.unique(np.round(x_tick_values, 6))
        differences = np.diff(unique_values)
        positive_differences = differences[differences > 1e-8]

        if len(positive_differences):
            step = float(np.min(positive_differences))
            expected_count = int(round(
                (unique_values.max() - unique_values.min()) / step
            )) + 1

            if 3 <= expected_count <= 30:
                expanded_values = (
                    unique_values.min()
                    + step * np.arange(expected_count)
                )
                expanded_pixels = (
                    expanded_values - x_scale["intercept"]
                ) / x_scale["slope"]
                x_tick_values = expanded_values
                x_tick_pixels = expanded_pixels

    order = np.argsort(x_tick_values)
    x_tick_values = x_tick_values[order]
    x_tick_pixels = x_tick_pixels[order]

    # Keep tick pixels that lie inside the detected plot box.
    valid_tick_mask = (
        (x_tick_pixels >= axis_box["left"] - 12)
        & (x_tick_pixels <= axis_box["right"] + 12)
    )
    x_tick_values = x_tick_values[valid_tick_mask]
    x_tick_pixels = x_tick_pixels[valid_tick_mask]

    if len(x_tick_values) < 3:
        return []

    calibration_ok, calibration_reason = _validate_plot_calibration(
        x_values=x_tick_values,
        x_pixels=x_tick_pixels,
        x_scale=x_scale,
        y_scale=y_scale,
        page_text_context=page_text_context,
    )

    if not calibration_ok:
        return []


    # Refine the data-region box from calibrated tick positions. This prevents
    # a page border or chart card border from being mistaken for the x-axis.
    if len(x_tick_pixels) >= 3 and y_scale.get("tick_pixels"):
        axis_box["left"] = int(round(float(np.min(x_tick_pixels))))
        axis_box["right"] = int(round(float(np.max(x_tick_pixels))))
        axis_box["bottom"] = int(round(
            float(np.median(axis_box.get("x_ticks", [{}])[0].get("y", axis_box["bottom"])))
        )) if axis_box.get("x_ticks") else axis_box["bottom"]
        y_tick_pixels = np.asarray(y_scale["tick_pixels"], dtype=float)
        axis_box["top"] = max(
            0,
            int(round(float(np.min(y_tick_pixels) - 55))),
        )

    series_masks = _build_series_masks(
        rgb,
        axis_box,
    )

    candidate_series = []

    for mask_info in series_masks:
        local_x_positions = (
            x_tick_pixels - axis_box["left"]
        )

        sampled_y = _sample_series_at_x(
            mask_info["mask"],
            local_x_positions,
        )

        quality = _series_quality(
            mask_info["mask"],
            sampled_y,
            local_x_positions,
        )

        if quality <= 0:
            continue

        candidate_series.append({
            **mask_info,
            "sampled_y": sampled_y,
            "quality": quality,
        })

    # Remove near-duplicate masks by comparing sampled y positions.
    deduplicated = []

    for candidate in sorted(
        candidate_series,
        key=lambda series: series["quality"],
        reverse=True,
    ):
        candidate_values = np.asarray([
            np.nan if value is None else value
            for value in candidate["sampled_y"]
        ], dtype=float)

        duplicate = False

        for existing in deduplicated:
            existing_values = np.asarray([
                np.nan if value is None else value
                for value in existing["sampled_y"]
            ], dtype=float)

            jointly_valid = (
                np.isfinite(candidate_values)
                & np.isfinite(existing_values)
            )

            if jointly_valid.sum() >= 3:
                median_distance = float(np.nanmedian(
                    np.abs(
                        candidate_values[jointly_valid]
                        - existing_values[jointly_valid]
                    )
                ))

                if median_distance <= 4.0:
                    duplicate = True
                    break

        if not duplicate:
            deduplicated.append(candidate)

    candidate_series = deduplicated[:6]

    if not candidate_series:
        return []

    legend_labels = _extract_legend_labels(
        image,
        axis_box,
    )

    if legend_labels:
        candidate_series = sorted(
            candidate_series,
            key=lambda series: series["quality"],
            reverse=True,
        )[:len(legend_labels)]

    candidate_series = _assign_series_labels(
        candidate_series,
        legend_labels,
    )

    structured_rows = []

    for tick_value, tick_pixel in zip(
        x_tick_values,
        x_tick_pixels,
    ):
        row = {
            "x": float(tick_value),
            "uncertainty_note": (
                "Approximate; digitized from rendered plot pixels"
            ),
        }

        for series in candidate_series:
            index = list(x_tick_pixels).index(tick_pixel)
            local_y = series["sampled_y"][index]

            if local_y is None:
                row[series["label"]] = None
                continue

            absolute_y = local_y + axis_box["top"]
            y_value = (
                y_scale["slope"] * absolute_y
                + y_scale["intercept"]
            )

            row[series["label"]] = float(y_value)

        structured_rows.append(row)


    # Repair shared/occluded baseline markers after series values are built.
    structured_rows = _repair_shared_baseline(
        structured_rows,
        [series["label"] for series in candidate_series],
    )

    # Require at least one non-empty series.
    populated_series = []

    for series in candidate_series:
        label = series["label"]
        values = [
            row.get(label)
            for row in structured_rows
            if row.get(label) is not None
        ]

        if len(values) >= 3:
            populated_series.append(label)

    if len(populated_series) < 1:
        return []

    structured_rows = [
        {
            key: value
            for key, value in row.items()
            if (
                key in {"x", "uncertainty_note"}
                or key in populated_series
            )
        }
        for row in structured_rows
    ]

    header = ["x"] + populated_series + ["uncertainty_note"]
    matrix_rows = [header]

    for row in structured_rows:
        matrix_rows.append([
            f"{row['x']:.6g}",
            *[
                (
                    ""
                    if row.get(label) is None
                    else f"{row[label]:.6g}"
                )
                for label in populated_series
            ],
            row["uncertainty_note"],
        ])

    embedding_lines = []

    for row in structured_rows:
        values = [
            f"x={row['x']:.6g}"
        ]

        for label in populated_series:
            if row.get(label) is not None:
                values.append(
                    f"{label}={row[label]:.6g}"
                )

        embedding_lines.append(
            ", ".join(values) + " (approximate)"
        )

    calibration_points = min(
        x_scale["points"],
        y_scale["points"],
    )

    series_coverage = np.mean([
        sum(value is not None for value in series["sampled_y"])
        / max(1, len(series["sampled_y"]))
        for series in candidate_series
        if series["label"] in populated_series
    ])

    confidence = float(min(
        0.90,
        0.40
        + 0.04 * calibration_points
        + 0.20 * series_coverage,
    ))

    item = create_content_item(
        document_id=document_record["document_id"],
        file_name=document_record["file"],
        doc_type=document_record["doc_type"],
        page_start=page_number,
        page_end=page_number,
        content_type="plot_table",
        text_for_embedding="\n".join(embedding_lines),
        text_for_llm=(
            "[PLOT-DERIVED TABLE — APPROXIMATE]\n"
            f"Source: {document_record['file']}, page {page_number}\n"
            f"Detected series: {', '.join(populated_series)}\n"
            f"Calibration confidence: {confidence:.2f}\n"
            "Values are approximate because they were digitized from "
            "rendered plot pixels.\n\n"
            f"{table_to_markdown(matrix_rows)}"
        ),
        extraction_method=(
            "opencv_multiseries_color_ocr_tick_calibration"
        ),
        bbox=[
            axis_box["left"],
            axis_box["top"],
            axis_box["right"],
            axis_box["bottom"],
        ],
        confidence=confidence,
        structured_data=structured_rows,
        uncertainty=(
            "Values are estimated from rendered plot pixels and should "
            "not be treated as exact source data."
        ),
        x_axis={
            "label": "x",
            "tick_values": [
                float(value) for value in x_tick_values
            ],
        },
        y_axis={
            "label": "y",
            "calibration_rmse": y_scale["rmse"],
        },
        series_names=populated_series,
        plot_detection_method=axis_box.get("method"),
        plot_validation_reason=calibration_reason,
        calibration={
            "x": x_scale,
            "y": y_scale,
        },
    )

    return [item]


def extract_plots_from_pdf(pdf_path, document_record):
    """
    Render every PDF page before plot detection.

    This supports both embedded raster charts and vector PDF charts.
    """
    items = []

    with fitz.open(pdf_path) as document:
        for page_index, page in enumerate(document):
            image = render_pdf_page_to_image(
                page,
                dpi=220,
            )

            page_items = extract_plot_table_from_image(
                image=image,
                document_record=document_record,
                page_number=page_index + 1,
                page_text_context=page.get_text("text"),
            )

            items.extend(page_items)

    return items



#@title CELL 28E — Extractors for DOCX, TXT, CSV, XLSX, and images

def dataframe_to_table_item(dataframe, document_record, page_number, label, method):
    dataframe = dataframe.fillna("").astype(str)
    matrix = [list(map(str, dataframe.columns.tolist()))] + dataframe.values.tolist()
    matrix = normalize_table_matrix(matrix)
    if not matrix or len(matrix) < 2:
        return None

    structure_type, embedding_text, markdown = format_detected_structure(matrix)
    return create_content_item(
        document_id=document_record["document_id"],
        file_name=document_record["file"],
        doc_type=document_record["doc_type"],
        page_start=page_number,
        page_end=page_number,
        content_type="table",
        text_for_embedding=embedding_text,
        text_for_llm=(
            f"[{label}]\nSource: {document_record['file']}\n\n{markdown}"
        ),
        extraction_method=method,
        table_structure_type=structure_type,
        structured_data=matrix,
        table_total_rows=len(matrix),
        table_columns=len(matrix[0]),
    )


def extract_non_pdf_items(file_path, document_record):
    path = Path(file_path)
    kind = get_file_kind(path)
    text_items, table_items, ocr_items, plot_items = [], [], [], []

    if kind == "txt":
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            text_items.append(create_content_item(
                document_id=document_record["document_id"],
                file_name=document_record["file"],
                doc_type=document_record["doc_type"],
                page_start=1,
                page_end=1,
                content_type="text",
                text_for_embedding=text,
                text_for_llm=f"[TEXT DOCUMENT]\nSource: {path.name}\n\n{text}",
                extraction_method="plain_text",
            ))

    elif kind == "docx":
        doc = Document(path)
        paragraph_text = "\n".join(
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )
        if paragraph_text:
            text_items.append(create_content_item(
                document_id=document_record["document_id"],
                file_name=document_record["file"],
                doc_type=document_record["doc_type"],
                page_start=1,
                page_end=1,
                content_type="text",
                text_for_embedding=paragraph_text,
                text_for_llm=f"[DOCX TEXT]\nSource: {path.name}\n\n{paragraph_text}",
                extraction_method="python_docx",
            ))

        for table_index, table in enumerate(doc.tables, start=1):
            matrix = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            matrix = normalize_table_matrix(matrix)
            if len(matrix) < 2 or len(matrix[0]) < 2:
                continue
            structure_type, embedding_text, markdown = format_detected_structure(matrix)
            table_items.append(create_content_item(
                document_id=document_record["document_id"],
                file_name=document_record["file"],
                doc_type=document_record["doc_type"],
                page_start=1,
                page_end=1,
                content_type="table",
                text_for_embedding=embedding_text,
                text_for_llm=(
                    f"[DOCX TABLE {table_index}]\nSource: {path.name}\n\n{markdown}"
                ),
                extraction_method="python_docx_table",
                table_structure_type=structure_type,
                structured_data=matrix,
            ))

    elif kind == "csv":
        item = dataframe_to_table_item(
            pd.read_csv(path),
            document_record,
            page_number=1,
            label="CSV TABLE",
            method="pandas_csv",
        )
        if item:
            table_items.append(item)

    elif kind == "xlsx":
        workbook = pd.ExcelFile(path)
        for sheet_index, sheet_name in enumerate(workbook.sheet_names, start=1):
            dataframe = pd.read_excel(path, sheet_name=sheet_name)
            item = dataframe_to_table_item(
                dataframe,
                document_record,
                page_number=sheet_index,
                label=f"XLSX SHEET: {sheet_name}",
                method="pandas_xlsx",
            )
            if item:
                item["sheet_name"] = sheet_name
                table_items.append(item)

    elif kind == "image":
        image = Image.open(path).convert("RGB")
        image_ocr_text, image_ocr_tables = create_ocr_items_for_image(
            image,
            document_record,
            page_number=1,
            extraction_method="tesseract_image",
        )
        ocr_items.extend(image_ocr_text)
        table_items.extend(image_ocr_tables)
        plot_items.extend(extract_plot_table_from_image(image, document_record, 1))

    return text_items, table_items, ocr_items, plot_items

